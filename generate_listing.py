import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================
# CONFIG (Расширения и игнорируемые папки)
# ==========================
ALLOWED_EXTENSIONS = {
    ".py", ".js", ".ts", ".html", ".css", ".json", ".md"
}

IGNORED_DIRS = {
    ".git", "__pycache__", "node_modules",
    "dist", "build", ".idea", ".vscode", ".venv"
}

OUTPUT_FILE = "report_2cols.docx"

# ==========================
# WORD HELPERS (Настройка колонок и шрифтов)
# ==========================
def set_two_columns(section):
    """Настраивает двухколоночный режим для секции Word"""
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    for child in sectPr.findall(qn('w:cols')):
        sectPr.remove(child)
    sectPr.append(cols)

def clean_text(text):
    """Убирает управляющие символы, которые могут сломать Word-парсер"""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)

def add_code(doc, text):
    """Добавляет блок кода шрифтом Courier New 8pt с нулевыми отступами"""
    text = clean_text(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Courier New')
    run.font.size = Pt(8)

# ==========================
# FILE HANDLING (Фильтрация файлов)
# ==========================
def should_include(file):
    """Проверяет, нужно ли обрабатывать файл по расширению"""
    if file.startswith("."):
        return False
    _, ext = os.path.splitext(file)
    return ext.lower() in ALLOWED_EXTENSIONS

def read_file(path):
    """Читает файл в кодировке UTF-8 с обходом ошибок чтения"""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return None

def get_description(filename, filepath):
    """Генерация описаний файлов для отчета интернет-магазина Fan Sport"""
    name = filename.lower()
    path = filepath.lower()
    
    if name == "manage.py":
        return "Точка входа Django проекта (запуск и управление сервером)"
    if name == "settings.py":
        return "Конфигурация проекта Django и настройки базы данных"
    if name == "urls.py":
        return "Настройка маршрутов и URL-адресов приложения"
    if name == "admin.py":
        return "Регистрация моделей в панели администратора Django Admin"
    if name == "apps.py":
        return "Конфигурация Django приложения"
    if name == "forms.py":
        return "Пользовательские формы регистрации и валидации"
    if name == "seed_products.py":
        return "Скрипт автоматического наполнения базы данных 20 спортивными товарами по категориям"
        
    # Специфические вьюхи и модели
    if name == "views.py":
        if "users" in path:
            return "Обработчики авторизации, регистрации и личного кабинета покупателя"
        return "Логика обработки каталога товаров, создания заказов и платежей"
        
    if name == "models.py":
        if "users" in path:
            return "Описание модели расширенного профиля пользователя"
        return "Описание моделей базы данных (товары, категории, заказы, отзывы)"
        
    # HTML Шаблоны
    if name == "landing.html":
        return "HTML шаблон главной страницы (каталог товаров с фильтрацией по категориям)"
    if name == "profile.html":
        return "HTML шаблон личного кабинета покупателя со списком оформленных заказов"
    if name == "payment_page.html":
        return "HTML шаблон интерактивного симулятора платежной системы ЮKassa"
    if name == "payment_success.html":
        return "HTML шаблон страницы успешного чека транзакции"
    if name == "base.html":
        return "Базовый мастер-шаблон разметки сайта со встроенной AJAX-корзиной"
    if name == "login.html":
        return "HTML шаблон страницы входа в аккаунт"
    if name == "register.html":
        return "HTML шаблон страницы регистрации нового пользователя"
        
    if name.endswith(".js"):
        return "Фронтенд логика (JavaScript)"
    if name.endswith(".html"):
        return "HTML шаблон страницы пользовательского интерфейса"
    if name.endswith(".css"):
        return "Стили оформления интерфейса"
        
    return "Исходный код программного модуля"

# ==========================
# MAIN GENERATOR LOGIC
# ==========================
def main():
    doc = Document()
    
    # Настраиваем первую секцию на две колонки
    section = doc.sections[0]
    set_two_columns(section)
    
    # Настраиваем узкие поля (~1.27 см) для экономии страниц
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    
    processed_files = 0
    
    # Рекурсивный обход директории
    for root, dirs, files in os.walk("."):
        # Исключаем скрытые и системные директории (окружение, кэш, гит)
        dirs[:] = [d for d in dirs if not d.startswith('.') and d not in IGNORED_DIRS]
        
        for file in files:
            if should_include(file):
                file_path = os.path.join(root, file)
                content = read_file(file_path)
                
                if content and content.strip():
                    processed_files += 1
                    lines_count = len(content.splitlines())
                    desc = get_description(file, file_path)
                    
                    # Название файла (жирным)
                    p_title = doc.add_paragraph()
                    p_title.paragraph_format.space_before = Pt(12)
                    p_title.paragraph_format.space_after = Pt(2)
                    p_title.paragraph_format.keep_with_next = True
                    
                    run_name = p_title.add_run(f"{file}\n")
                    run_name.bold = True
                    run_name.font.size = Pt(11)
                    run_name.font.name = 'Arial'
                    
                    # Описание файла (курсивом)
                    run_desc = p_title.add_run(f"Описание: {desc}\n")
                    run_desc.italic = True
                    run_desc.font.size = Pt(9)
                    run_desc.font.name = 'Arial'
                    
                    # Количество строк
                    run_lines = p_title.add_run(f"Lines: {lines_count}")
                    run_lines.font.size = Pt(8.5)
                    run_lines.font.name = 'Arial'
                    
                    # Код файла
                    add_code(doc, content)
                    
                    # Короткий визуальный разделитель
                    p_sep = doc.add_paragraph()
                    p_sep.paragraph_format.space_before = Pt(4)
                    p_sep.paragraph_format.space_after = Pt(4)
                    run_sep = p_sep.add_run("~" * 35)
                    run_sep.font.size = Pt(8)
                    run_sep.font.name = 'Courier New'
                    
                    print(f"Обработан файл: {file_path} ({lines_count} строк)")
    
    if processed_files > 0:
        doc.save(OUTPUT_FILE)
        print(f"\nУспешно! Двухколоночный листинг сохранен в файл: {OUTPUT_FILE}")
    else:
        print("\nФайлы исходного кода для обработки не найдены.")

if __name__ == "__main__":
    main()